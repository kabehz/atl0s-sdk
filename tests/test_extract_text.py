from docx import Document
from PIL import Image, ImageDraw
from pypdf import PdfWriter
import pytesseract
import unittest
import sys
from pathlib import Path

# Añadir el directorio raíz del proyecto al PYTHONPATH
sys.path.append(str(Path(__file__).resolve().parent.parent))

from validador_service_v4 import extract_text

class TestExtractText(unittest.TestCase):
    def setUp(self):
        self.anexos_dir = Path(__file__).parent / "anexos"
        self.anexos_dir.mkdir(exist_ok=True)

        # Archivos creados durante las pruebas
        self.created_files = []

        self.pdf_file = self.anexos_dir / "sample.pdf"
        self.docx_file = self.anexos_dir / "sample.docx"
        self.image_file = self.anexos_dir / "sample.png"
        self.unsupported_file = self.anexos_dir / "sample.xlsx"

        # # Crear un archivo PDF válido con texto real
        # writer = PdfWriter()
        # writer.add_blank_page(width=72, height=72)
        # with open(self.pdf_file, "wb") as f:
        #     writer.write(f)
        # self.created_files.append(self.pdf_file)

        # # Crear un archivo DOCX válido
        # doc = Document()
        # doc.add_paragraph("Dummy DOCX content")
        # doc.save(self.docx_file)
        # self.created_files.append(self.docx_file)

        # # Crear una imagen válida
        # img = Image.new("RGB", (100, 100), color="white")
        # draw = ImageDraw.Draw(img)
        # draw.text((10, 10), "Dummy Image content", fill="black")
        # img.save(self.image_file)
        # self.created_files.append(self.image_file)

        # # Crear un archivo no soportado
        # self.unsupported_file.write_text("Dummy Unsupported content")
        # self.created_files.append(self.unsupported_file)

    def tearDown(self):
        # Eliminar solo los archivos creados durante las pruebas
        for file in self.created_files:
            file.unlink(missing_ok=True)



# class TestExtractText(unittest.TestCase):
#     def setUp(self):
#         self.anexos_dir = Path(__file__).parent / "anexos"
#         self.anexos_dir.mkdir(exist_ok=True)

#         self.pdf_file = self.anexos_dir / "sample.pdf"
#         self.docx_file = self.anexos_dir / "sample.docx"
#         self.image_file = self.anexos_dir / "sample.png"
#         self.unsupported_file = self.anexos_dir / "sample.xlsx"

#         # Crear un archivo PDF válido con texto real
#         writer = PdfWriter()
#         writer.add_blank_page(width=72, height=72)
#         with open(self.pdf_file, "wb") as f:
#             writer.write(f)

#         # Crear un archivo DOCX válido
#         doc = Document()
#         doc.add_paragraph("Dummy DOCX content")
#         doc.save(self.docx_file)

#         # Crear una imagen válida
# #        img = Image.new("RGB", (100, 100), color="white")
# #        draw = ImageDraw.Draw(img)
# #        draw.text((10, 10), "Dummy Image content", fill="black")
# #        img.save(self.image_file)

#         # Crear un archivo no soportado
#         self.unsupported_file.write_text("Dummy Unsupported content")

#     def tearDown(self):
#         # Eliminar los archivos creados
#         for file in [self.pdf_file, self.docx_file, self.image_file, self.unsupported_file]:
#             file.unlink(missing_ok=True)



#     # def test_extract_text_image(self):
#     #     # Crear una imagen válida para la prueba
#     #     img = Image.new("RGB", (200, 100), color="white")  # Aumentar el tamaño para mejorar la precisión del OCR
#     #     draw = ImageDraw.Draw(img)
#     #     draw.text((10, 40), "Dummy Image content", fill="black")  # Ajustar posición y tamaño del texto
#     #     img.save(self.image_file)

#     #     # Extraer texto de la imagen
#     #     text = extract_text(self.image_file)

#     #     # Verificar que el texto esperado esté contenido en el texto extraído
#     #     self.assertIn("Dummy Image content", text)
    
#     def test_extract_text_image(self):
#         # Crear una imagen válida para la prueba
#         img = Image.new("RGB", (300, 150), color="white")  # Aumentar el tamaño de la imagen
#         draw = ImageDraw.Draw(img)

#         # Usar una fuente clara y ajustar el tamaño del texto
#         try:
#             from PIL import ImageFont
#             font = ImageFont.truetype("arial.ttf", 24)  # Usar una fuente estándar
#         except IOError:
#             font = None  # Usar la fuente predeterminada si no está disponible

#         draw.text((10, 50), "Dummy Image content", fill="black", font=font)  # Ajustar posición y tamaño del texto
#         img.save(self.image_file)

#         # Extraer texto de la imagen
#         text = extract_text(self.image_file)

#         # Verificar que el texto esperado esté contenido en el texto extraído
#         self.assertIn("Dummy Image content", text)

#     def test_extract_text_unsupported_file_type(self):
#         with self.assertRaises(ValueError):
#             extract_text(self.unsupported_file)


if __name__ == "__main__":
    unittest.main()
